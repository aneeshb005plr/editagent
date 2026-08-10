"""
app/documents/docx_parser.py

Parses .docx files into the common ParsedDocument representation
(app/documents/base.py). Built and verified against python-docx 1.2.0
- see inline notes for API specifics confirmed against the real
library rather than assumed from memory (document.part.related_parts
is a dict, NOT a related_part() method - that would have been a
real bug if guessed).

KNOWN SIMPLIFICATION, previously present, now RESOLVED: paragraph-vs-
table interleaving order was not preserved by the old doc.paragraphs
+ doc.tables approach. This is now fixed via doc.iter_inner_content()
(and _Cell.iter_inner_content() for nested content), confirmed to
exist and yield correct document order in python-docx 1.2.0. This
same refactor also fixes nested tables (see below) - both were the
same underlying limitation of the old flat-loop approach.

NESTED TABLES ARE NOW HANDLED: a table inside a table cell is
invisible to doc.tables entirely (python-docx's own documentation
states only top-level tables appear there) - confirmed by direct
reproduction, where a nested table's content did not appear via the
old doc.tables-based approach at all. _process_table now recurses
into any table found via cell.iter_inner_content(), at any nesting
depth, using the same table_counter so nested tables get their own
distinct table_index rather than colliding with their parent's.

NOT YET LOAD-TESTED AT 100MB - python-docx loads the full document
into memory; this is the parser most in need of the large-file
spike flagged in the architecture doc.

EXPLICITLY OUT OF SCOPE, not silently missing: headers, footers,
footnotes/endnotes, and textboxes live in separate document parts
that doc.paragraphs/doc.tables never touch - none of this content
is reviewed OR flagged as unsupported; it's simply not visited at
all. Stated here explicitly rather than left implicit, since
"readable content" could otherwise be assumed to mean everything in
the file. Revisit if review coverage for these sections becomes a
real requirement - each is a genuinely separate python-docx access
path (doc.sections[i].header/.footer, etc.), not a small extension
of the current paragraph walk.

ALSO EXPLICITLY OUT OF SCOPE, confirmed by direct test: text marked
as a tracked-change DELETION (w:del/w:delText) is excluded from
paragraph.text - confirmed by injecting a real w:del element and
observing it does not appear in the parsed text. A document with
unaccepted deletions will not have that deleted text reviewed, and
it is not flagged as unsupported either (it's simply invisible to
python-docx's own .text property, the same as headers/footers).
Revisit only if reviewing tracked-change content becomes a real
requirement - would need direct XML inspection for w:ins/w:del,
not something paragraph.text can be coaxed into providing.

KNOWN SIMPLIFICATION: table column_index is a cell ordinal from
enumerate(row.cells), not a true grid position - merged/spanned
cells can make the reported column_index inaccurate relative to the
visual grid. Acceptable for MVP; a finding's presence and row are
still correct, only exact column alignment under merges may drift.
"""

from __future__ import annotations

import io
import logging

from docx import Document as DocxDocument
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)
from docx.oxml.ns import nsmap, qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.documents.base import (
    ContentBlock,
    ContentKind,
    Location,
    ParsedDocument,
    UnsupportedItem,
    UnsupportedKind,
    ExtractionStatus,
)

logger = logging.getLogger("app.documents.docx_parser")

_WORD_DRAWING_NS = (
    "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
)
_INLINE_TAG = f"{_WORD_DRAWING_NS}inline"
_ANCHOR_TAG = f"{_WORD_DRAWING_NS}anchor"
_GRAPHIC_DATA_TAG = (
    "{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData"
)

# Sourced from python-docx's OWN nsmap rather than hardcoded strings -
# these values were verified to already be correct by direct
# comparison, but importing from the library's source of truth
# prevents any future drift, per review feedback.
_PICTURE_URI = nsmap["pic"]
_CHART_URI = nsmap["c"]
_DIAGRAM_URI = nsmap["dgm"]

# Object/OLE embed tags live in a different namespace - checked broadly
# via localname to avoid a second namespace-map maintenance burden.
_OLE_OBJECT_LOCALNAMES = {"object", "OLEObject"}


def _is_heading(paragraph: Paragraph) -> bool:
    # Confirmed: python-docx exposes the applied style name directly;
    # built-in heading styles are named "Heading 1".."Heading 9".
    style_name = paragraph.style.name if paragraph.style else ""
    return style_name.lower().startswith("heading")


def _extract_picture(
    graphic_el,
    doc: DocxDocument,
    location_for_log: str,
) -> UnsupportedItem:
    """Resolves an inline PICTURE to its raw bytes via the document
    part's relationship map.

    location_for_log is a plain display string for log messages ONLY
    (e.g. "paragraph 3" or "table 0 cell (0,1)") - NOT used to build
    the actual UnsupportedItem's Location, which the caller sets
    correctly afterward. Previously this parameter was a bare int
    always labeled "paragraph %d" in log messages, so a warning about
    a table-cell picture would misleadingly print "paragraph 3" when
    the real location was table 3 - correct final data, wrong
    diagnostics, confirmed by inspection.

    Confirmed API (python-docx 1.2.0): doc.part.related_parts is a
    dict keyed by relationship id -> Part. There is no
    related_part(rId) method on DocumentPart in this version.

    Handles two real cases beyond the straightforward embedded-raster
    path, both confirmed rather than assumed:
    - LINKED (not embedded) pictures use r:link instead of r:embed -
      blip.embed is None for these, which previously got misreported
      as a generic parse failure. Now correctly identified and
      flagged as linked rather than broken.
    - EMF/WMF and other vector/unusual image formats: image_part.
      image.blob raises UnrecognizedImageError/InvalidImageStreamError
      /UnexpectedEndOfFileError from docx.image.exceptions - confirmed
      these are plain Exception subclasses, NOT AttributeError, so a
      bare `except AttributeError` does not catch them and the whole
      parse would crash on a real, common case (e.g. a pasted Excel
      chart or Visio object saved as EMF). Falls back to the part's
      raw .blob (bytes without python-docx's own format parsing/
      validation) - confirmed correct attribute name by inspecting
      opc.part.Part.blob's actual source, not just reasoned about.
    """

    # Navigate via raw XML find(), NOT the .graphic convenience
    # attribute - confirmed by testing that CT_Anchor (floating/
    # wrapped images) does not expose .graphic the way CT_Inline
    # does (AttributeError: 'CT_Anchor' object has no attribute
    # 'graphic'). This approach works identically for both.
    a_ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    blip = graphic_el.find(f".//{a_ns}blip")

    if blip is None:
        logger.warning(
            "Graphic at %s has no blip element - flagging without "
            "extraction",
            location_for_log,
        )
        return UnsupportedItem(
            kind=UnsupportedKind.IMAGE,
            location=Location(),  # caller overwrites
            note="unparseable inline picture structure",
            extraction_status=ExtractionStatus.FAILED,
        )

    r_id = blip.get(qn("r:embed"))

    if r_id is None:
        link_r_id = blip.get(qn("r:link"))
        if link_r_id is not None:
            return UnsupportedItem(
                kind=UnsupportedKind.IMAGE,
                location=Location(),  # caller overwrites
                note="linked (not embedded) image - no local bytes to extract",
                extraction_status=ExtractionStatus.NOT_APPLICABLE,
            )
        logger.warning(
            "Inline picture at %s has neither r:embed nor r:link - "
            "flagging without extraction",
            location_for_log,
        )
        return UnsupportedItem(
            kind=UnsupportedKind.IMAGE,
            location=Location(),  # caller overwrites
            note="unparseable inline picture structure",
            extraction_status=ExtractionStatus.FAILED,
        )

    image_part = doc.part.related_parts.get(r_id)

    if image_part is None:
        logger.warning(
            "Could not resolve part for rId=%s at %s",
            r_id,
            location_for_log,
        )
        return UnsupportedItem(
            kind=UnsupportedKind.IMAGE,
            location=Location(),  # caller overwrites
            note=f"could not resolve part for rId={r_id}",
            extraction_status=ExtractionStatus.FAILED,
        )

    try:
        raw_bytes = image_part.image.blob
        content_type = image_part.content_type
    except (
        UnrecognizedImageError,
        InvalidImageStreamError,
        UnexpectedEndOfFileError,
    ):
        # python-docx's own image-format parsing failed (EMF/WMF and
        # similar are common real-world triggers) - fall back to the
        # part's raw bytes, which exist regardless of whether python-
        # docx can parse the format's header. content_type is a plain
        # always-present attribute on a Part (confirmed via source),
        # so no getattr guard needed here.
        logger.info(
            "Image at %s is in a format python-docx can't parse "
            "(likely EMF/WMF) - using raw part bytes instead",
            location_for_log,
        )
        raw_bytes = image_part.blob
        content_type = image_part.content_type

    return UnsupportedItem(
        kind=UnsupportedKind.IMAGE,
        location=Location(),  # caller overwrites
        note=f"content_type={content_type}",
        raw_bytes=raw_bytes,
        extraction_status=ExtractionStatus.PENDING,
    )


def _extract_inline_graphics(
    paragraph: Paragraph,
    doc: DocxDocument,
    location_for_log: str,
) -> list[UnsupportedItem]:
    """Finds every graphic (picture, chart, or SmartArt) inside this
    paragraph - both INLINE (wp:inline, flows with text) and FLOATING/
    WRAPPED (wp:anchor, positioned with text wrap) - and routes each
    to the right handling by inspecting graphicData's uri. Searching
    only wp:inline was a real, silent coverage gap: any image with
    text wrapping applied (a common real-world formatting choice, not
    an edge case) uses wp:anchor instead and was previously invisible
    to this parser entirely - not flagged, not extracted, just absent.

    Returned items carry a PLACEHOLDER Location() - the caller (either
    the body-paragraph loop or the table-cell loop in parse_docx)
    always sets the real .location afterward. This is deliberate,
    not an oversight: the previous version built the Location here
    using a bare int meaning "paragraph index" in the body-paragraph
    call site but was passed "table index" in the cell call site
    (with the caller overwriting .location afterward regardless) -
    correct final data, but any warning logged INSIDE this function
    during the call would print "paragraph 3" for what was actually
    table 3. Now nothing is built here that isn't immediately
    overwritten, and location_for_log is a plain, honest display
    string used ONLY for log messages.

    Charts/SmartArt in Word have no python-docx API support at all
    (confirmed: no docx.chart module exists), so - matching
    pptx_parser's approach - they're flagged only, never label-
    extracted the way PowerPoint charts are; that PPT-specific
    capability has no Word equivalent with the current library."""

    items: list[UnsupportedItem] = []
    graphic_elements = paragraph._p.findall(f".//{_INLINE_TAG}") + paragraph._p.findall(
        f".//{_ANCHOR_TAG}"
    )

    # Defensive dedup by embed rId: mc:AlternateContent wraps the
    # same drawing in both mc:Choice and mc:Fallback for backward-
    # compatibility with older Word versions - a real OOXML pattern
    # real Word-authored documents can contain (not something python-
    # docx itself generates, but something it can open). An unbounded
    # `.//` search finds both copies, which would otherwise double-
    # extract the same image and double-count its bytes/findings.
    seen_r_ids: set[str] = set()

    for graphic_el in graphic_elements:
        graphic_data = graphic_el.find(f".//{_GRAPHIC_DATA_TAG}")
        uri = graphic_data.get("uri", "") if graphic_data is not None else ""

        if uri == _PICTURE_URI:
            a_ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
            blip_el = graphic_el.find(f".//{a_ns}blip")
            r_id = blip_el.get(qn("r:embed")) if blip_el is not None else None

            if r_id is not None:
                if r_id in seen_r_ids:
                    continue  # AlternateContent duplicate - skip
                seen_r_ids.add(r_id)

            items.append(_extract_picture(graphic_el, doc, location_for_log))

        elif uri == _CHART_URI:
            items.append(
                UnsupportedItem(
                    kind=UnsupportedKind.CHART,
                    location=Location(),  # caller overwrites
                    note=(
                        "chart data/visual layout not reviewed - "
                        "Word charts have no label-extraction path "
                        "(unlike PowerPoint) with the current library"
                    ),
                    extraction_status=ExtractionStatus.NOT_APPLICABLE,
                )
            )

        elif uri == _DIAGRAM_URI:
            items.append(
                UnsupportedItem(
                    kind=UnsupportedKind.SMARTART,
                    location=Location(),  # caller overwrites
                    note="SmartArt diagram - structure/content not reviewed in current scope",
                    extraction_status=ExtractionStatus.NOT_APPLICABLE,
                )
            )

        else:
            # NOTE: python-docx's own InlineShape.type classification
            # uses this exact same URI-first check (confirmed by
            # reading its source) and would ALSO return
            # NOT_IMPLEMENTED for a wrapped graphic (e.g. Word Drawing
            # Canvas / wordprocessingCanvas) whose graphicData uri
            # isn't pic/chart/diagram even though a real pic:pic sits
            # nested inside. This isn't a divergence from the
            # library's own behavior - but probing for a nested blip
            # anyway, beyond what the library itself does, costs
            # little and can recover real image bytes the library's
            # own classification would also miss.
            a_ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
            nested_blip = graphic_el.find(f".//{a_ns}blip")
            if nested_blip is not None and nested_blip.get(qn("r:embed")):
                logger.info(
                    "Graphic at %s has unrecognized graphicData "
                    "uri=%r but contains a nested blip - attempting "
                    "extraction anyway",
                    location_for_log,
                    uri,
                )
                items.append(_extract_picture(graphic_el, doc, location_for_log))
                continue

            logger.warning(
                "Graphic at %s has unrecognized graphicData uri=%r "
                "and no extractable blip - flagging as unsupported",
                location_for_log,
                uri,
            )
            items.append(
                UnsupportedItem(
                    kind=UnsupportedKind.EMBEDDED_OBJECT,
                    location=Location(),  # caller overwrites
                    note=f"unrecognized graphic type (uri={uri or 'none'})",
                    extraction_status=ExtractionStatus.NOT_APPLICABLE,
                )
            )

    return items


def _has_ole_object(paragraph: Paragraph) -> bool:
    for elem in paragraph._p.iter():
        tag_localname = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag_localname in _OLE_OBJECT_LOCALNAMES:
            return True
    return False


def _process_paragraph(
    paragraph: Paragraph,
    location: Location,
    doc: DocxDocument,
    parsed: ParsedDocument,
    kind_if_text: ContentKind,
) -> None:
    """Shared logic for handling one paragraph's text + graphics + OLE
    at an already-determined location. Used identically for body-level
    paragraphs and paragraphs inside table cells (including nested
    table cells) - the only difference between call sites is what
    Location and ContentKind get passed in."""

    text = paragraph.text.strip()
    if text:
        parsed.blocks.append(
            ContentBlock(text=text, kind=kind_if_text, location=location)
        )

    graphics = _extract_inline_graphics(paragraph, doc, location.display())
    for item in graphics:
        item.location = location
    parsed.unsupported_items.extend(graphics)

    if _has_ole_object(paragraph):
        parsed.unsupported_items.append(
            UnsupportedItem(
                kind=UnsupportedKind.EMBEDDED_OBJECT,
                location=location,
                note="OLE-embedded object - no extraction path in current scope",
                extraction_status=ExtractionStatus.NOT_APPLICABLE,
            )
        )


def _process_table(
    table: Table,
    doc: DocxDocument,
    parsed: ParsedDocument,
    table_counter: list[int],
) -> None:
    """Processes a table's cells, RECURSING into any nested tables
    found via cell.iter_inner_content() - a table inside a cell is
    invisible to doc.tables entirely (confirmed: python-docx's own
    docs state only top-level tables appear there), and was
    completely unhandled before this refactor - no block, no
    unsupported item, silently absent. table_counter is shared and
    incremented for every table encountered at ANY nesting level, so
    nested tables get their own distinct table_index rather than
    colliding with their parent's."""

    this_table_index = table_counter[0]
    table_counter[0] += 1

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_location = Location(
                table_index=this_table_index,
                row_index=row_idx,
                column_index=col_idx,
            )

            for cell_item in cell.iter_inner_content():
                if isinstance(cell_item, Paragraph):
                    _process_paragraph(
                        cell_item,
                        cell_location,
                        doc,
                        parsed,
                        ContentKind.TABLE_CELL,
                    )
                else:
                    # A Table nested inside this cell - recurse using
                    # the exact same machinery, at any depth.
                    _process_table(cell_item, doc, parsed, table_counter)


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parses a .docx file's readable content into a ParsedDocument.

    File-size validation is the DISPATCHER's responsibility (single
    source of truth against settings.MAX_FILE_SIZE_MB) - this
    function assumes it's already been called with an acceptable
    size and focuses purely on content extraction.
    """

    doc = DocxDocument(io.BytesIO(file_bytes))
    parsed = ParsedDocument(source_filename=filename, file_type="docx")

    paragraph_counter = [0]
    table_counter = [0]

    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            idx = paragraph_counter[0]
            paragraph_counter[0] += 1
            location = Location(paragraph_index=idx)
            kind = (
                ContentKind.HEADING if _is_heading(item) else ContentKind.PARAGRAPH
            )
            _process_paragraph(item, location, doc, parsed, kind)
        else:
            _process_table(item, doc, parsed, table_counter)

    logger.info(
        "Parsed %s: %d text blocks, %d unsupported items (%.0f chars)",
        filename,
        len(parsed.blocks),
        len(parsed.unsupported_items),
        parsed.total_char_count,
    )

    return parsed